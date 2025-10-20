#!/usr/bin/env python3
"""
Exporterar alla artiklar till Word-format (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import glob

def create_word_document():
    """Skapar ett Word-dokument med alla artiklar"""
    doc = Document()
    
    # Sätt upp dokumentstil
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Titel
    title = doc.add_heading('Flex Applications - Kunskapsbas', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Intro
    intro = doc.add_paragraph()
    intro.add_run('Komplett samling av alla artiklar från knowledge.flexapplications.se\n').bold = True
    intro.add_run(f'Exporterad: 2025-10-20\n')
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Tom rad
    
    # Kategorier
    categories = {
        'systemgemensamt': 'Systemgemensamt',
        'time': 'Time',
        'employee': 'Employee',
        'travel-expense': 'Travel & Expense',
        'payroll': 'Payroll',
        'plan': 'Plan'
    }
    
    article_count = 0
    
    for category_id, category_name in categories.items():
        # Läs alla markdown-filer i kategorin
        category_path = f'documentation/{category_id}'
        if not os.path.exists(category_path):
            continue
        
        md_files = sorted(glob.glob(f'{category_path}/*.md'))
        
        if not md_files:
            continue
        
        # Kategori-rubrik
        doc.add_page_break()
        category_heading = doc.add_heading(f'{category_name}', 1)
        category_heading_run = category_heading.runs[0]
        category_heading_run.font.color.rgb = RGBColor(79, 70, 229)  # Indigo färg
        
        doc.add_paragraph(f'{len(md_files)} artiklar')
        doc.add_paragraph()  # Tom rad
        
        # Läs och lägg till varje artikel
        for md_file in md_files:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extrahera titel och innehåll
            lines = content.strip().split('\n')
            title = lines[0].replace('# ', '') if lines else 'Utan titel'
            
            # Artikel-rubrik
            article_heading = doc.add_heading(title, 2)
            
            # Artikelinnehåll
            article_text = '\n'.join(lines[1:]).strip()
            
            # Ta bort markdown-formatering för bättre läsbarhet i Word
            article_text = article_text.replace('## ', '')
            article_text = article_text.replace('### ', '')
            article_text = article_text.replace('**', '')
            article_text = article_text.replace('*', '')
            
            # Lägg till innehåll
            if article_text:
                paragraphs = article_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            doc.add_paragraph()  # Tom rad mellan artiklar
            article_count += 1
    
    # Spara dokumentet
    output_file = 'alla_artiklar.docx'
    doc.save(output_file)
    print(f'✅ Word-dokument skapat: {output_file}')
    print(f'📊 Totalt {article_count} artiklar exporterade')
    
    # Visa filstorlek
    file_size = os.path.getsize(output_file)
    size_mb = file_size / (1024 * 1024)
    print(f'📦 Filstorlek: {size_mb:.1f} MB')

if __name__ == '__main__':
    create_word_document()

