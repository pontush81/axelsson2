#!/usr/bin/env python3
"""
Exporterar alla artiklar till olika format MED bilder
- Word (.docx) med inbäddade bilder
- Excel (.xlsx) med länkar till bilder
- ZIP med alla filer
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, Alignment
import os
import json
import shutil
from pathlib import Path
from datetime import datetime

def export_to_word_with_images():
    """Exporterar till Word med inbäddade bilder"""
    print("\n📄 Skapar Word-dokument med bilder...")
    
    doc = Document()
    
    # Titel
    title = doc.add_heading('Flex HRM Dokumentation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Info
    info = doc.add_paragraph()
    info.add_run(f'Exporterad: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').bold = True
    info.add_run('Inkluderar text och bilder\n')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    
    categories = {
        'systemgemensamt': 'Systemgemensamt',
        'time': 'Time',
        'employee': 'Employee',
        'travel-expense': 'Travel & Expense',
        'payroll': 'Payroll',
        'plan': 'Plan'
    }
    
    total_articles = 0
    total_images = 0
    
    for category_id, category_name in categories.items():
        index_file = f'documentation/{category_id}/index.json'
        
        if not os.path.exists(index_file):
            continue
        
        # Läs index
        with open(index_file, 'r', encoding='utf-8') as f:
            articles = json.load(f)
        
        # Kategori-rubrik
        doc.add_heading(f'{category_name} ({len(articles)} artiklar)', 1)
        
        for article in articles:
            # Artikel-titel
            doc.add_heading(article['title'], 2)
            
            # Metadata
            meta = doc.add_paragraph()
            meta.add_run(f"📅 {article['date']}\n").italic = True
            meta.add_run(f"🔗 {article['url']}\n").italic = True
            if article.get('imageCount', 0) > 0:
                meta.add_run(f"📸 {article['imageCount']} bilder\n").italic = True
            
            # Läs artikel-fil
            article_file = f"documentation/{category_id}/{article['file']}"
            if os.path.exists(article_file):
                with open(article_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Ta bort metadata-sektion (allt innan ----)
                if '----' in content:
                    content = content.split('----', 1)[1].strip()
                
                # Lägg till text
                doc.add_paragraph(content[:1000] + '...' if len(content) > 1000 else content)
            
            # Lägg till bilder om de finns
            if article.get('hasImages'):
                img_folder = f"documentation/{category_id}/images"
                if os.path.exists(img_folder):
                    # Hitta bilder för denna artikel
                    slug = article['slug']
                    article_images = [f for f in os.listdir(img_folder) if slug in f]
                    
                    for img_file in article_images[:5]:  # Max 5 bilder per artikel
                        img_path = f"{img_folder}/{img_file}"
                        try:
                            doc.add_picture(img_path, width=Inches(4))
                            total_images += 1
                        except Exception as e:
                            print(f"  ⚠️  Kunde inte lägga till bild {img_file}: {e}")
            
            doc.add_paragraph()  # Mellanrum
            total_articles += 1
        
        doc.add_page_break()
    
    # Spara
    output_file = f'flex_hrm_dokumentation_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(output_file)
    
    print(f"✅ Word-dokument skapat: {output_file}")
    print(f"   📝 {total_articles} artiklar")
    print(f"   📸 {total_images} bilder inbäddade")
    
    return output_file

def export_to_excel():
    """Exporterar till Excel med länkar"""
    print("\n📊 Skapar Excel-fil...")
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Flex HRM Artiklar"
    
    # Headers
    headers = ['Kategori', 'Titel', 'Datum', 'URL', 'Innehåll (förkortad)', 'Antal bilder', 'Bildmapp']
    ws.append(headers)
    
    # Stil för header
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
        cell.alignment = Alignment(horizontal="center")
    
    categories = {
        'systemgemensamt': 'Systemgemensamt',
        'time': 'Time',
        'employee': 'Employee',
        'travel-expense': 'Travel & Expense',
        'payroll': 'Payroll',
        'plan': 'Plan'
    }
    
    total_articles = 0
    
    for category_id, category_name in categories.items():
        index_file = f'documentation/{category_id}/index.json'
        
        if not os.path.exists(index_file):
            continue
        
        with open(index_file, 'r', encoding='utf-8') as f:
            articles = json.load(f)
        
        for article in articles:
            # Förkortad innehåll
            excerpt = article.get('excerpt', '')[:200]
            
            # Bildmapp
            img_folder = f"documentation/{category_id}/images" if article.get('hasImages') else ''
            
            ws.append([
                category_name,
                article['title'],
                article['date'],
                article['url'],
                excerpt,
                article.get('imageCount', 0),
                img_folder
            ])
            
            total_articles += 1
    
    # Auto-anpassa kolumnbredder
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Spara
    output_file = f'flex_hrm_dokumentation_{datetime.now().strftime("%Y%m%d")}.xlsx'
    wb.save(output_file)
    
    print(f"✅ Excel-fil skapad: {output_file}")
    print(f"   📝 {total_articles} artiklar")
    
    return output_file

def create_export_zip():
    """Skapar ZIP med allt (artiklar + bilder)"""
    print("\n📦 Skapar ZIP-arkiv med allt innehåll...")
    
    zip_name = f'flex_hrm_complete_{datetime.now().strftime("%Y%m%d")}'
    
    # Skapa temp-mapp
    temp_dir = 'temp_export'
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    # Kopiera hela documentation-mappen
    shutil.copytree('documentation', f'{temp_dir}/documentation')
    
    # Skapa ZIP
    shutil.make_archive(zip_name, 'zip', temp_dir)
    
    # Cleanup
    shutil.rmtree(temp_dir)
    
    file_size = os.path.getsize(f'{zip_name}.zip') / (1024 * 1024)
    
    print(f"✅ ZIP skapad: {zip_name}.zip")
    print(f"   📦 Storlek: {file_size:.1f} MB")
    print(f"   📁 Innehåller: Alla artiklar + alla bilder")
    
    return f'{zip_name}.zip'

def main():
    """Kör alla exporter"""
    print("="*70)
    print("  FLEX HRM DOKUMENTATION - EXPORT MED BILDER")
    print("="*70)
    
    try:
        # Word med bilder
        word_file = export_to_word_with_images()
        
        # Excel
        excel_file = export_to_excel()
        
        # ZIP med allt
        zip_file = create_export_zip()
        
        print("\n" + "="*70)
        print("✅ EXPORT KLAR!")
        print("="*70)
        print(f"\n📄 Word: {word_file}")
        print(f"📊 Excel: {excel_file}")
        print(f"📦 ZIP (komplett): {zip_file}")
        print("\n💡 Tips:")
        print("  - Word: Öppna direkt i Microsoft Word eller Google Docs")
        print("  - Excel: Öppna i Excel/Numbers/Google Sheets")
        print("  - ZIP: Extrahera för tillgång till alla filer och bilder")
        print("="*70 + "\n")
        
    except Exception as e:
        print(f"\n❌ Fel vid export: {e}")
        print("Installera dependencies: pip install python-docx openpyxl")

if __name__ == '__main__':
    main()

